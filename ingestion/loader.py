"""
loader.py — Load raw documents from disk.
Supports: PDF, DOCX, TXT, Markdown.
Returns a list of Document(page_content, metadata) objects.
"""
import os
from pathlib import Path
from dataclasses import dataclass, field
from typing import List

@dataclass
class Document:
    page_content: str
    metadata: dict = field(default_factory=dict)


def load_pdf(path: str) -> List[Document]:
    from pypdf import PdfReader
    reader = PdfReader(path)
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": path, "page": i + 1, "type": "pdf"}
            ))
    return docs


def load_docx(path: str) -> List[Document]:
    from docx import Document as DocxDocument
    doc = DocxDocument(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(page_content=text, metadata={"source": path, "type": "docx"})]


def load_text(path: str) -> List[Document]:
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()
    ext = Path(path).suffix.lower()
    return [Document(page_content=text, metadata={"source": path, "type": ext.lstrip(".")})]


LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_text,
    ".md": load_text,
}


def load_directory(dir_path: str) -> List[Document]:
    """Recursively load all supported documents from a directory."""
    docs = []
    for root, _, files in os.walk(dir_path):
        for fname in files:
            ext = Path(fname).suffix.lower()
            if ext in LOADERS:
                full_path = os.path.join(root, fname)
                try:
                    loaded = LOADERS[ext](full_path)
                    docs.extend(loaded)
                    print(f"  ✓ Loaded {full_path} ({len(loaded)} segments)")
                except Exception as e:
                    print(f"  ✗ Failed to load {full_path}: {e}")
    return docs
